from flask import Flask, jsonify, request
import psycopg2
import json
from docx import Document
from fpdf import FPDF

# Connect to PostgreSQL
conn = psycopg2.connect(
    dbname="poc", user="postgres", password="xFl201X20FabSN2024Nov", host="localhost", port="5432"
)

# Fetch Speaker Data 
def fetch_speaker_data():
    with conn.cursor() as cursor:
        cursor.execute("SELECT * FROM Speakers")
        return cursor.fetchall()
    
# Fetch Agenda Data 
def fetch_agenda_data():
    
    query = """
    SELECT * 
    FROM agenda;
    """
    with conn.cursor() as cursorII:
        cursorII.execute(query)
        return cursorII.fetchall()

#Generate Json data
def generate_json(data):
    return json.dumps(data, indent=4)

#Generate Json DataII
def generate_jsonII(agenda_data):
    formatted_data = []
    for row in agenda_data:
        session = {
            "time": row[0],
            "type": row[1],
            "title": row[2],
            "speaker": row[3],
            "description": row[4],
        }
        formatted_data.append(session)
    with open("agenda.json", "w") as json_file:
        json.dump({"agenda": formatted_data}, json_file, indent=4)

speakerData = fetch_speaker_data()
print(speakerData)
print("Erics speaker data \n")
json_serializable_data = speakerData.tobytes()
jsondata = json.dumps({"data": json_serializable_data.decode("utf-8")})
##jsondata = generate_json(speakerData)
print(jsondata) 
print("Erics speaker jason data \n")


agendaData=fetch_agenda_data()
print(agendaData)
jsondataII= generate_json (agendaData)   #going to have to figure out date times... for now lets finalize the database
print(jsondataII)


conn.close()



# Example usage: more info in safari gpt


def generate_pdf(data, file_name):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, str(data))
    pdf.output(file_name)

def generate_doc(data, file_name):
    doc = Document()
    doc.add_heading('Your Document Title', level=1)
    doc.add_paragraph(str(data))
    doc.save(file_name)

###gpt on how to save  the file locally

