import docx 
import tempfile
import os
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

class WordDocument:
    def __init__(self, filename):
        self.doc = docx.Document()
        self.filename = filename

    def create_heading(self, heading_text, heading_level: int): #should be 1-6
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(heading_text)
        font = run.font
        font.name = 'Arial'
        font.bold = True
        font.size = Pt(24)  # Increased font size
        font.color.rgb = RGBColor(0x34, 0x6a, 0xc5)  # Blue color
        paragraph.style = self.doc.styles['Heading {}'.format(heading_level)]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.space_before = Pt(12)  # Add some space before the heading
        paragraph.space_after = Pt(6)  # Add some space after the heading

        
    def add_session(self, title, time, abstract):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(f'{time} : {title}')
        font = run.font
        font.name = 'Arial'
        font.bold = True
        font.size = Pt(16)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run(abstract)
        
        
    def add_image(self, image_data):
        if image_data:
            # Create a temporary file to store the image
            with tempfile.NamedTemporaryFile(suffix='.jpg') as tmp_file:
                # Write the BYTEA data to the temporary file
                tmp_file.write(image_data)
                tmp_file.flush()

                # Add the image to the document
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run()
                run.add_picture(tmp_file.name, width=Inches(2.5))

                # Remove the temporary file
                os.remove(tmp_file.name)
    
    def add_image2(self, image_data):
        if image_data:
            # Create a temporary file to store the image
            tmp_file = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
            # Write the BYTEA data  to the temporary file
            tmp_file.write(image_data)
            tmp_file.close()

            # Add the image to the document
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run()
            run.add_picture(tmp_file.name, width=Inches(2.5))

            # Remove the temporary file
            os.remove(tmp_file.name)

    def add_event_details(self, date, description):
        # Add the date in large orange font
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(date)
        font = run.font
        font.name = 'Arial'
        font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)  # Orange color
        font.size = Pt(24)

        # Add the description in large blue font
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(description)
        font = run.font
        font.name = 'Arial'
        #font.color.theme_color = WD_COLOR_INDEX.TEXT2
        font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue color
        font.size = Pt(24)
        
    def add_speaker_list(self, speaker_name, title, company):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(f'{speaker_name}, {title} at ')
        font = run.font
        font.name = 'Arial'
        font.size = Pt(12)

        run = paragraph.add_run(company)
        font = run.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = True
        
    def speakers_to_doc(self, speaker_processed):
        for i in range(0, len(speaker_processed), 3):
            name = speaker_processed[i]
            title = speaker_processed[i+ 1]
            company = speaker_processed[i +2]     
            self.add_speaker_list(name, title, company)
            
            
            
  
   
    def create_total_heading(self, heading_text, image_data): #should be 1-6
        if image_data:
            # Create a temporary file to store the image
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp_file:
                # Write the BYTEA data to the temporary file
                tmp_file.write(image_data)
                tmp_file.flush()

                # Add the image to the document
                table = self.doc.add_table(rows=1, cols=2, style='Table Grid')
                cell = table.cell(0, 0)
                run = cell.add_paragraph().add_run()
                run.add_picture(tmp_file.name, width=Inches(1.5))

                # Remove the temporary file
                os.remove(tmp_file.name)
            table = self.doc.tables[-1]
            cell = table.cell(0, 1)
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(heading_text)
            font = run.font
            font.name = 'Arial'
            font.bold = True
            font.size = Pt(24)  # Increased font size
            font.color.rgb = RGBColor(0x34, 0x6a, 0xc5)  # Blue color
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.space_before = Pt(12)  # Add some space before the heading
            paragraph.space_after = Pt(6)  # Add some space after the heading

    def save(self):
        self.doc.save(self.filename)